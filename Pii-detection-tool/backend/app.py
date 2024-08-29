from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import spacy
from presidio_analyzer import AnalyzerEngine
import docx
import pdfplumber
import os

app = Flask(__name__)
CORS(app)  # Enables CORS for all routes

# Load SpaCy model
nlp = spacy.load("en_core_web_sm")

# Initialize the Presidio Analyzer Engine
analyzer = AnalyzerEngine()

# Define the 18 types of PII to detect
pii_types = ["PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER", "IP_ADDRESS", "NRIC", "DATE_TIME", "CREDIT_CARD", 
             "IBAN_CODE", "US_PASSPORT", "US_DRIVER_LICENSE", "SSN", "LOCATION", "ORG", "GPE", "MONEY", 
             "ID", "AADHAAR_NUMBER", "LICENSE_PLATE"]

def detect_pii(text):
    """Detect PII in the provided text."""
    results = analyzer.analyze(text=text, language="en", entities=pii_types)
    return results

def list_detected_pii(text, results):
    """List detected PII types and their values."""
    pii_list = {}
    for result in results:
        entity_text = text[result.start:result.end]
        pii_list[entity_text] = result.entity_type  # Use actual text as the key
    return pii_list

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"})

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"})

    if file and (file.filename.endswith('.txt') or file.filename.endswith('.docx') or file.filename.endswith('.pdf')):
        file_path = os.path.join('uploads', file.filename)
        file.save(file_path)

        if file.filename.endswith('.txt'):
            with open(file_path, 'r') as f:
                text = f.read()
        elif file.filename.endswith('.docx'):
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text = '\n'.join(full_text)
        elif file.filename.endswith('.pdf'):
            with pdfplumber.open(file_path) as pdf:
                full_text = ""
                for page in pdf.pages:
                    full_text += page.extract_text() + "\n"
            text = full_text

        pii_results = detect_pii(text)
        pii_list = list_detected_pii(text, pii_results)

        return jsonify({"pii_list": pii_list, "file_path": file_path})
    else:
        return jsonify({"error": "Unsupported file format"})

@app.route('/mask', methods=['POST'])
def mask_file():
    data = request.json
    file_path = data['file_path']
    edited_pii_data = data['edited_pii_data']

    if not os.path.isfile(file_path):
        return jsonify({"error": "File not found"})

    if file_path.endswith('.txt'):
        with open(file_path, 'r') as f:
            text = f.read()
    elif file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text = '\n'.join(full_text)
    elif file_path.endswith('.pdf'):
        with pdfplumber.open(file_path) as pdf:
            full_text = ""
            for page in pdf.pages:
                full_text += page.extract_text() + "\n"
        text = full_text

    # Apply edits from edited_pii_data
    for entity_text, new_value in edited_pii_data.items():
        text = text.replace(entity_text, new_value)

    masked_file_path = file_path.replace('.txt', '_masked.txt')
    with open(masked_file_path, 'w') as f:
        f.write(text)

    return jsonify({"masked_file_path": masked_file_path})

@app.route('/download/<path:filename>', methods=['GET'])
def download_file(filename):
    return send_file(filename, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
